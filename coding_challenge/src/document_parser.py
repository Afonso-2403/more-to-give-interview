from pathlib import Path

import openpyxl
from docx import Document

from models import Foundation


def parse_project_description(path: str) -> str:
    """Read a .docx file and return all text (paragraphs + tables) as a single string."""
    doc = Document(path)

    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def parse_text_file(path: str) -> str:
    """Read a plain .txt file and return its content."""
    return Path(path).read_text(encoding="utf-8")


def load_project_description(path: str) -> str:
    """Dispatch to the right parser based on file extension."""
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return parse_project_description(path)
    elif ext == ".txt":
        return parse_text_file(path)
    else:
        raise ValueError(f"Unsupported project file type: {ext}. Use .docx or .txt")


def parse_foundations_list(path: str) -> list[Foundation]:
    """Read an .xlsx file and return a list of Foundation objects.

    Expects columns: (number, name, url) with a header row.
    Skips rows where URL is empty.
    """
    wb = openpyxl.load_workbook(path, read_only=True)
    ws = wb.active

    foundations: list[Foundation] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i == 0:  # skip header
            continue
        number, name, url = row[0], row[1], row[2]
        if not url or not name:
            continue
        foundations.append(Foundation(number=int(number), name=str(name), url=str(url)))

    wb.close()
    return foundations
